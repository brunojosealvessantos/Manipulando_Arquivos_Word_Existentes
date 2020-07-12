from docx import Document

meu_word = Document('desafio.docx')
meu_word.add_paragraph("Saudações, Mestres da Autonação!!!", style='Title')

paragrafo = meu_word.add_paragraph('Seja bem vindo ao curso')
paragrafo.add_run('Mestres automaçâo', style='Strong')
paragrafo.add_run('um lugar de aprendizado')

meu_word.add_paragraph("Módulo 1 - Mestre do Software", style='Heading 1')
meu_word.add_paragraph(
    "Módulo 2 - Mestre do Bootcamp Python", style='Heading 2')
meu_word.add_paragraph("Módulo 3 - Mestre da Web", style='Heading 3')
meu_word.add_paragraph("Descubra o poder da automação", style='Quote')
meu_word.add_paragraph('Variaveis', style='List Bullet')
meu_word.add_paragraph('Condicionais', style='List Bullet')
meu_word.add_paragraph('Repetição', style='List Bullet')
meu_word.add_paragraph('Classe', style='List Bullet')
meu_word.add_paragraph('Herança', style='List Bullet')
meu_word.add_paragraph('Word', style='List Number')
meu_word.add_paragraph('Excel', style='List Bullet')
meu_word.add_paragraph('Power Point', style='List Bullet')

aulas = {
    (1, 'Instalando Python', '3:07'),
    (2, 'Instalando Selenium', '7:34'),
    (1, 'Instalando openpyxl', '2:45')
}

tabela_aulas = meu_word.add_table(rows=1, cols=3)

cabecalho_tabela = tabela_aulas.rows[0].cells
cabecalho_tabela[0].text = 'Aula'
cabecalho_tabela[1].text = 'Nome'
cabecalho_tabela[2].text = 'Duração'

for aula, nome, duracao in aulas:
    dados_por_linha = tabela_aulas.add_row().cells
    dados_por_linha[0].text = str(aula)
    dados_por_linha[1].text = nome
    dados_por_linha[2].text = duracao

meu_word.add_picture('foto.jpg')


meu_word.save('desafio.docx')
